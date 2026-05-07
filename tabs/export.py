"""
Tab: Exportação
Oferece opções de exportação (PDF, Word, WhatsApp) e compartilhamento.
"""
import logging
from io import BytesIO
from pathlib import Path
from urllib.parse import quote

import streamlit as st

from clause_service import extrair_pontos_atencao

logger = logging.getLogger(__name__)


def _gerar_pdf(analise_texto: str, nome_arquivo: str = "analise_contrato") -> BytesIO:
    from clause_service import extrair_numeros_riscos
    from report_service import gerar_relatorio_pdf

    riscos = extrair_numeros_riscos(analise_texto)
    risk_counts = {"altos": riscos["altos"], "medios": riscos["medios"], "baixos": riscos["baixos"]}

    return gerar_relatorio_pdf(
        analysis_text=analise_texto,
        contract_name=nome_arquivo,
        risk_counts=risk_counts,
    )


def _gerar_word(analise_texto: str, nome_arquivo: str = "analise_contrato") -> BytesIO:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    try:
        doc = Document()

        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        titulo = doc.add_heading("ANÁLISE DE CONTRATO", level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitulo = doc.add_paragraph(nome_arquivo)
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitulo_format = subtitulo.runs[0]
        subtitulo_format.italic = True
        subtitulo_format.font.size = Pt(9)

        doc.add_paragraph()

        for linha in analise_texto.split("\n"):
            if not linha.strip():
                continue

            if linha.startswith("# "):
                p = doc.add_heading(linha.replace("# ", ""), level=1)
            elif linha.startswith("## "):
                p = doc.add_heading(linha.replace("## ", ""), level=2)
            elif linha.startswith("### "):
                p = doc.add_heading(linha.replace("### ", ""), level=3)
            else:
                p = doc.add_paragraph(linha)

                for run in p.runs:
                    if any(
                        word in run.text
                        for word in [
                            "RISCO ALTO",
                            "RISCO MÉDIO",
                            "Cláusula:",
                            "Problema:",
                            "Base legal:",
                            "Sugestão",
                        ]
                    ):
                        run.bold = True

        doc.add_paragraph()
        rodape = doc.add_paragraph(
            "Documento gerado pela plataforma TRUST CORPORATION - Contrato Seguro IA"
        )
        for run in rodape.runs:
            run.font.size = Pt(8)
            run.italic = True

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    except Exception as e:
        logger.error(f"Erro ao gerar Word: {str(e)}")
        raise


def render_export():
    """Renderiza a aba de exportação com opções PDF, Word e WhatsApp."""
    st.markdown(
        '<h3 style="color: #3A6FA0;">💾 Exportar e Compartilhar</h3>',
        unsafe_allow_html=True,
    )
    st.write("Baixe a análise completa ou compartilhe o resumo com seus clientes.")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown(
            '<div class="info-card" style="text-align: center;">',
            unsafe_allow_html=True,
        )
        st.markdown("#### PDF Oficial")
        st.write("Documento formatado ideal para impressão e envio formal.")
        if st.button("📄 Gerar PDF", use_container_width=True):
            try:
                with st.spinner("Gerando PDF..."):
                    pdf_buffer = _gerar_pdf(
                        st.session_state.analise_resultado,
                        st.session_state.nome_arquivo,
                    )
                    st.download_button(
                        label="⬇️ Baixar PDF",
                        data=pdf_buffer,
                        file_name=f"analise_{Path(st.session_state.nome_arquivo).stem}.pdf",
                        mime="application/pdf",
                        use_container_width=True,
                        key="btn_dl_pdf",
                    )
                    st.success("Pronto!")
            except Exception as e:
                st.error(f"Erro ao gerar PDF: {str(e)}")
        st.markdown("</div>", unsafe_allow_html=True)

    with col2:
        st.markdown(
            '<div class="info-card" style="text-align: center;">',
            unsafe_allow_html=True,
        )
        st.markdown("#### Word (Editável)")
        st.write(
            "Documento editável (.docx) para você adicionar seus próprios comentários."
        )
        if st.button("📝 Gerar Word", use_container_width=True):
            try:
                with st.spinner("Gerando Word..."):
                    word_buffer = _gerar_word(
                        st.session_state.analise_resultado,
                        st.session_state.nome_arquivo,
                    )
                    st.download_button(
                        label="⬇️ Baixar Word",
                        data=word_buffer,
                        file_name=f"analise_{Path(st.session_state.nome_arquivo).stem}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                        key="btn_dl_word",
                    )
                    st.success("Pronto!")
            except Exception as e:
                st.error(f"Erro ao gerar Word: {str(e)}")
        st.markdown("</div>", unsafe_allow_html=True)

    with col3:
        st.markdown(
            '<div class="info-card" style="text-align: center;">',
            unsafe_allow_html=True,
        )
        st.markdown("#### 📱 WhatsApp (Twilio)")
        st.write("PDF + resumo enviados automaticamente.")

        telefone = st.text_input(
            "Seu número de WhatsApp:",
            placeholder="5514999999999",
            value=st.session_state.get("user_whatsapp", ""),
            key="whatsapp_twilio_phone",
        )

        if st.button("🚀 Enviar PDF + Resumo", use_container_width=True, type="primary"):
            if not telefone:
                st.warning("Informe seu número de WhatsApp")
            else:
                try:
                    with st.spinner("Gerando PDF e enviando via Twilio..."):
                        from notification_service import NotificationService

                        notifier = NotificationService()
                        pdf_buf = _gerar_pdf(
                            st.session_state.analise_resultado,
                            st.session_state.nome_arquivo,
                        )
                        resumo = extrair_pontos_atencao(
                            st.session_state.analise_resultado, max_linhas=20
                        )
                        phone = "".join(c for c in telefone if c.isdigit())
                        fname = f"analise_{Path(st.session_state.nome_arquivo).stem}.pdf"

                        ok, msg = notifier.send_whatsapp_pdf_twilio(
                            to_number=phone,
                            pdf_buffer=pdf_buf,
                            filename=fname,
                            summary=resumo,
                            contract_name=st.session_state.nome_arquivo,
                        )

                        if ok:
                            st.success(f"Enviado com sucesso para {phone}!")
                        else:
                            st.error(f"Falha: {msg}")
                            st.info("Verifique se o número está autorizado no sandbox Twilio (join parts-current)")

                except Exception as e:
                    st.error(f"Erro: {str(e)}")

        with st.expander("📖 Autorize o envio do WhatsApp", expanded=False):
            st.markdown("""
            **Para poder receber os arquivos, faça somente 1 vez:**

            Envie uma mensagem de seu número de WhatsApp para o número:
            ```
            +1 415 523 8886
            ```
            Com a mensagem:
            ```
            join parts-current
            ```
            """)

        st.markdown("---")

        st.markdown("#### 🔗 Link manual (fallback)")
        st.write("Se o Twilio falhar, envie o resumo manualmente.")
        if st.button("📱 Abrir WhatsApp com resumo", use_container_width=True):
            resumo = extrair_pontos_atencao(
                st.session_state.analise_resultado, max_linhas=20
            )
            mensagem = (
                "📊 *ANÁLISE DE CONTRATO - TRUST CORPORATION*\n"
                f"📄 *Contrato:* {st.session_state.nome_arquivo}\n\n"
                "⚠️ *RESUMO DE RISCOS:*\n"
                f"{resumo}\n\n"
                "✅ TRUST CORPORATION - Contrato Seguro IA"
            )
            st.markdown(
                f'<a href="https://api.whatsapp.com/send?text={quote(mensagem[:4000])}" target="_blank">'
                '<button style="width: 100%; padding: 0.75rem; background-color: #25D366; '
                'color: white; border: none; border-radius: 0.4rem; cursor: pointer; '
                'font-weight: bold;">📱 Abrir WhatsApp</button></a>',
                unsafe_allow_html=True,
            )

        st.markdown("</div>", unsafe_allow_html=True)
