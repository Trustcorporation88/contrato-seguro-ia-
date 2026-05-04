"""
extractor_service.py - Serviço de Extração Multi-Formato

Extrai texto de diversos formatos de documento: DOCX, RTF, ODT, imagens (OCR).
"""

import logging
from io import BytesIO
from pathlib import Path
from typing import Optional, Tuple

logger = logging.getLogger(__name__)

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    from pdf_extractor import extrair_texto_pdf_bytes

    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


ACCEPTED_EXTENSIONS = {
    ".pdf": "PDF",
    ".txt": "Texto",
    ".docx": "Word",
    ".doc": "Word (legado)",
    ".rtf": "Rich Text",
    ".odt": "OpenDocument",
    ".png": "Imagem",
    ".jpg": "Imagem",
    ".jpeg": "Imagem",
    ".tiff": "Imagem",
    ".bmp": "Imagem",
}


def detectar_formato(filename: str) -> str:
    """
    Detecta o formato do arquivo pela extensão.

    Returns:
        Nome do formato ou 'Desconhecido'
    """
    ext = Path(filename).suffix.lower()
    return ACCEPTED_EXTENSIONS.get(ext, "Desconhecido")


def extrair_texto_docx(file_bytes: BytesIO) -> str:
    """
    Extrai texto de um arquivo .docx.

    Args:
        file_bytes: BytesIO com o conteúdo do arquivo

    Returns:
        Texto extraído
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx não está instalado. Execute: pip install python-docx")

    file_bytes.seek(0)
    doc = Document(file_bytes)
    paragrafos = []

    for para in doc.paragraphs:
        if para.text.strip():
            paragrafos.append(para.text)

    # Extrair também de tabelas
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                paragrafos.append(" | ".join(row_text))

    return "\n".join(paragrafos)


def extrair_texto_imagem(file_bytes: BytesIO, lang: str = "por") -> str:
    """
    Extrai texto de uma imagem usando OCR.

    Args:
        file_bytes: BytesIO com a imagem
        lang: Idioma para OCR (padrão: por = português)

    Returns:
        Texto extraído
    """
    if not OCR_AVAILABLE:
        raise ImportError("pytesseract/Pillow não instalados. Execute: pip install pytesseract Pillow")

    file_bytes.seek(0)
    img = Image.open(file_bytes)
    texto = pytesseract.image_to_string(img, lang=lang)
    return texto.strip()


def extrair_texto_arquivo(
    file_bytes: BytesIO, filename: str, enable_ocr: bool = True
) -> Tuple[str, str]:
    """
    Extrai texto de qualquer formato suportado.

    Args:
        file_bytes: BytesIO com o conteúdo do arquivo
        filename: Nome original do arquivo (para detectar formato)
        enable_ocr: Para PDFs/imagens, usa OCR se necessário

    Returns:
        Tuple (texto_extraido, formato_detectado)

    Raises:
        ValueError: Formato não suportado
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf" and PDF_AVAILABLE:
        texto = extrair_texto_pdf_bytes(file_bytes, enable_ocr=enable_ocr)
        return texto, "PDF"

    elif ext == ".txt":
        file_bytes.seek(0)
        try:
            texto = file_bytes.read().decode("utf-8")
        except UnicodeDecodeError:
            file_bytes.seek(0)
            texto = file_bytes.read().decode("latin-1")
        return texto, "Texto"

    elif ext == ".docx":
        texto = extrair_texto_docx(file_bytes)
        return texto, "Word (DOCX)"

    elif ext in (".png", ".jpg", ".jpeg", ".tiff", ".bmp"):
        texto = extrair_texto_imagem(file_bytes)
        return texto, f"Imagem ({ext.upper()})"

    elif ext == ".rtf":
        file_bytes.seek(0)
        texto = file_bytes.read().decode("utf-8", errors="ignore")
        # RTF simples: remove tags de formatação
        import re

        texto = re.sub(r"\\[a-z]+\d*\s?", " ", texto)
        texto = re.sub(r"[\\{}]", "", texto)
        return texto.strip(), "RTF"

    elif ext == ".odt":
        file_bytes.seek(0)
        import zipfile
        from xml.etree import ElementTree

        with zipfile.ZipFile(file_bytes) as zf:
            if "content.xml" in zf.namelist():
                xml_content = zf.read("content.xml")
                root = ElementTree.fromstring(xml_content)
                ns = {"text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0"}
                textos = root.findall(".//text:p", ns)
                paragrafos = [t.text or "" for t in textos if t.text]
                return "\n".join(paragrafos), "ODT"

    return "", ACCEPTED_EXTENSIONS.get(ext, "Desconhecido")


def formatos_aceitos() -> list:
    """Retorna lista de extensões aceitas."""
    return list(ACCEPTED_EXTENSIONS.keys())
